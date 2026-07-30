#!/usr/bin/env python3
"""
보안 매거진 콘텐츠 산출물 생성기 (로컬 macOS 전용)

콘텐츠 JSON 한 개를 입력받아 다음을 생성한다.
  1. 보안콘텐츠_YYYYMMDD.docx  (서론/본문/결론 라벨 유지)
  2. 보안콘텐츠_YYYYMMDD.pdf   (서론+본문+결론 통합, 라벨 없음)
  3. magazine_list.csv 에 새 행 추가 (--csv 지정 시)

사용법(권장 — 겹치지 않는 폴더 자동 생성):
  python3 generate.py content.json --base "매거진 컨텐츠" \
      [--csv "magazine_list.csv"] [--tags "태그1, 태그2"] [--link "url"]

  → 실행일 기준 "매거진 컨텐츠/YYMMDD" 폴더를 만든다.
    같은 날 이미 있으면 "YYMMDD_2", "YYMMDD_3" ... 로 겹치지 않게 증가시킨다.

사용법(폴더를 직접 지정):
  python3 generate.py content.json --outdir "매거진 컨텐츠/260702" [...]

content.json 스키마는 references/content-schema.json 참고.
"""
import argparse, csv, json, os, shutil, sys
from datetime import date

# ── 한글 폰트 자동 탐색 ────────────────────────────────────────────
FONT_CANDIDATES = [
    "/Library/Fonts/NanumGothic.ttf",
    os.path.expanduser("~/Library/Fonts/NanumGothic.ttf"),
    "/System/Library/Fonts/Supplemental/AppleGothic.ttf",  # macOS 기본(단일 TTF)
]
DOCX_FONT = "Apple SD Gothic Neo"  # Word/뷰어가 이름으로 해석


def find_ttf():
    for p in FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    sys.exit("한글 TTF 폰트를 찾지 못했습니다. NanumGothic.ttf 를 ~/Library/Fonts 에 설치하세요.")


def resolve_outdir(base, yymmdd):
    """base/YYMMDD 폴더를 만든다. 이미 있으면 _2, _3 ... 으로 겹치지 않게 생성."""
    cand = os.path.join(base, yymmdd)
    if not os.path.exists(cand):
        os.makedirs(cand)
        return cand
    n = 2
    while True:
        cand = os.path.join(base, f"{yymmdd}_{n}")
        if not os.path.exists(cand):
            os.makedirs(cand)
            return cand
        n += 1


def load_content(path):
    with open(path, encoding="utf-8") as f:
        c = json.load(f)
    required = ["category", "title", "intro", "body_paragraphs", "action_guide", "conclusion"]
    missing = [k for k in required if not c.get(k)]
    if missing:
        sys.exit(f"content.json 필수 항목 누락: {missing}")
    c.setdefault(
        "notice",
        "※ 본 콘텐츠는 국내 보안뉴스 및 공개된 보안 권고를 참고하여 생성형 AI를 활용해 작성되었으며, 관리자 검수 후 게시되었습니다.",
    )
    c.setdefault("date", date.today().strftime("%Y%m%d"))
    return c


# ── DOCX 생성 (라벨 유지) ──────────────────────────────────────────
def make_docx(c, outdir):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)

    def para(text, size=10, bold=False, color=None, space_after=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = DOCX_FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    para("보안 매거진 콘텐츠", size=16, bold=True, space_after=2)
    para(f"{c['date']}  |  v4.01_local", size=9, color=(0x88, 0x88, 0x88), space_after=10)
    para(c["category"], size=10, color=(0x7C, 0x6F, 0xE0))
    para(c["title"], size=15, bold=True, space_after=10)

    para("[서론]", size=9, bold=True, color=(0x66, 0x66, 0x66), space_after=2)
    para(c["intro"])
    para("[본문]", size=9, bold=True, color=(0x66, 0x66, 0x66), space_after=2)
    for bp in c["body_paragraphs"]:
        para(bp)
    for item in c["action_guide"]:
        para(item, space_after=2)
    para("[결론]", size=9, bold=True, color=(0x66, 0x66, 0x66), space_after=2)
    para(c["conclusion"], space_after=12)

    para(c["notice"], size=8, color=(0x88, 0x88, 0x88))

    path = os.path.join(outdir, f"보안콘텐츠_{c['date']}.docx")
    doc.save(path)
    return path


# ── PDF 생성 (통합 본문) ───────────────────────────────────────────
def make_pdf(c, outdir):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    ttf = find_ttf()
    pdfmetrics.registerFont(TTFont("KR", ttf))

    path = os.path.join(outdir, f"보안콘텐츠_{c['date']}.pdf")
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=25 * mm, rightMargin=25 * mm, topMargin=25 * mm, bottomMargin=25 * mm,
    )
    s_cat = ParagraphStyle("cat", fontName="KR", fontSize=10, textColor="#7C6FE0")
    s_title = ParagraphStyle("title", fontName="KR", fontSize=16, leading=22, spaceAfter=8)
    s_body = ParagraphStyle("body", fontName="KR", fontSize=10, leading=18, spaceAfter=6)
    s_guide = ParagraphStyle("guide", fontName="KR", fontSize=10, leading=18, leftIndent=12)
    s_notice = ParagraphStyle("notice", fontName="KR", fontSize=8, textColor="#888888")

    story = [Paragraph(c["category"], s_cat), Spacer(1, 4 * mm),
             Paragraph(c["title"], s_title),
             HRFlowable(width="100%", thickness=1, color="#DDDDDD", spaceAfter=6),
             Paragraph(c["intro"], s_body), Spacer(1, 3 * mm)]
    for bp in c["body_paragraphs"]:
        story += [Paragraph(bp, s_body), Spacer(1, 3 * mm)]
    for item in c["action_guide"]:
        story.append(Paragraph(item, s_guide))
    story += [Spacer(1, 3 * mm), Paragraph(c["conclusion"], s_body), Spacer(1, 6 * mm),
              HRFlowable(width="100%", thickness=0.5, color="#CCCCCC", spaceAfter=4),
              Paragraph(c["notice"], s_notice)]
    doc.build(story)
    return path


# ── CSV 업데이트 ───────────────────────────────────────────────────
FIELDS = ["NO", "발행일", "카테고리", "매거진 제목", "핵심 키워드 (태그)", "상태", "비고 / 뉴스 링크"]


def update_csv(c, csv_path, tags, link):
    rows = []
    if os.path.exists(csv_path):
        with open(csv_path, encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))
    pub_date = f"{c['date'][:4]}-{c['date'][4:6]}-{c['date'][6:]}"
    # 카테고리에서 "1) 스미싱/피싱" → "스미싱/피싱" 만 추출
    cat = c["category"].split(")", 1)[-1].strip() if ")" in c["category"] else c["category"]
    rows.append({
        "NO": len(rows) + 1, "발행일": pub_date, "카테고리": cat,
        "매거진 제목": c["title"], "핵심 키워드 (태그)": tags,
        "상태": "작성 완료", "비고 / 뉴스 링크": link,
    })
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=FIELDS)
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k, "") for k in FIELDS})
    return csv_path, len(rows)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("content")
    ap.add_argument("--base", help="산출 폴더의 상위 경로. YYMMDD 하위 폴더를 겹치지 않게 자동 생성")
    ap.add_argument("--outdir", help="산출 폴더를 직접 지정 (--base 와 택일)")
    ap.add_argument("--csv")
    ap.add_argument("--tags", default="")
    ap.add_argument("--link", default="")
    ap.add_argument("--snapshot", action="store_true",
                    help="업데이트된 CSV 사본을 산출 폴더에 magazine_list_YYYYMMDD.csv 로 저장")
    a = ap.parse_args()

    if not a.base and not a.outdir:
        sys.exit("--base 또는 --outdir 중 하나는 반드시 지정해야 합니다.")

    c = load_content(a.content)

    if a.base:
        outdir = resolve_outdir(a.base, c["date"][2:])  # YYYYMMDD → YYMMDD
    else:
        outdir = a.outdir
        os.makedirs(outdir, exist_ok=True)
    print(f"📁 폴더: {outdir}")

    docx_path = make_docx(c, outdir)
    print(f"✅ DOCX: {docx_path}")
    pdf_path = make_pdf(c, outdir)
    print(f"✅ PDF : {pdf_path}")
    if a.csv:
        p, n = update_csv(c, a.csv, a.tags, a.link)
        print(f"✅ CSV : {p} ({n}건)")
        if a.snapshot:
            snap = os.path.join(outdir, f"magazine_list_{c['date']}.csv")
            shutil.copyfile(p, snap)
            print(f"✅ CSV 사본: {snap}")
    elif a.snapshot:
        print("⚠️  --snapshot 은 --csv 와 함께 사용해야 합니다. 사본을 만들지 않았습니다.")
    print(f"\n다음 배너 생성 시 이 폴더를 사용하세요: {outdir}")


if __name__ == "__main__":
    main()
